from __future__ import annotations

import csv
import html
import re
from pathlib import Path

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
DOC_DIR = OUTPUT / "doc"
PDF_DIR = OUTPUT / "pdf"
SOURCE_MD = OUTPUT / "sierra-tech-spaces-customer-offer-egyptian-arabic.md"
PROSPECT_MD = OUTPUT / "potential-customers-to-call-egypt.md"
PROSPECT_CSV = OUTPUT / "potential-customers-to-call-egypt.csv"
DOCX_OUTPUT = DOC_DIR / "sierra-tech-spaces-customer-offer-egyptian-arabic.docx"
PDF_OUTPUT = PDF_DIR / "sierra-tech-spaces-customer-offer-egyptian-arabic.pdf"
LOGO = ROOT / "STS_Logo_Gold_Outlined_Transparent.png"

FONT_REG = r"C:\Windows\Fonts\tahoma.ttf"
FONT_BOLD = r"C:\Windows\Fonts\tahomabd.ttf"

NAVY = colors.HexColor("#0B1426")
SLATE = colors.HexColor("#1E293B")
TEAL = colors.HexColor("#06B6D4")
GOLD = colors.HexColor("#D9B45F")
MUTED = colors.HexColor("#64748B")
LIGHT_BG = colors.HexColor("#F7FAFC")
RULE = colors.HexColor("#D7DEE8")


OFFER_SECTIONS = [
    {
        "title": "الوعد الأساسي",
        "body": [
            "Sierra Tech Spaces بتساعد الشركات تشيل الشغل المتكرر، ترد أسرع على العملاء، وتلم فرص بيع أكتر باستخدام أنظمة ذكاء اصطناعي عملية مبنية على مشاكلهم الحقيقية.",
            "احنا مش بنبدأ ببيع مشروع كبير. بنبدأ بمكالمة نفهم فيها فين الوقت والفلوس والفرص بتضيع، وبعدها نعمل ديمو صغير مجاني يورّي العميل الحل هيشتغل ازاي قبل ما يدفع.",
        ],
        "quote": "قولنا ببساطة: احكيلنا ايه اللي مبطأ الشغل عندك، واحنا هنرسم المشكلة، نعمل ديمو شغال، ونورّيك ايه ممكن يتعمله أتمتة قبل أي التزام.",
    },
    {
        "title": "بنقدملك ايه؟",
        "items": [
            {
                "heading": "1. جلسة اكتشاف المشكلة - مجانا",
                "bullets": [
                    "ايه أكتر مهام بتاكل وقت كل أسبوع؟",
                    "فين الـ leads بتضيع أو محدش بيرد عليها بسرعة؟",
                    "ايه الأسئلة اللي العملاء بيسألوها كل يوم؟",
                    "ايه العمليات اللي لسه ماشية بالمجهود اليدوي أو ذاكرة الموظفين؟",
                    "ايه التقارير أو الحجوزات أو الفواتير أو المتابعات اللي لسه بتتعمل يدوي؟",
                ],
                "outcome": "العميل يطلع فاهم يبتدي automation منين وأسرع مكسب هييجي منين.",
            },
            {
                "heading": "2. ديمو مخصوص - مجانا",
                "bullets": [
                    "مساعد واتساب بيرد على أسئلة حقيقية.",
                    "تأهيل leads من Facebook أو Instagram أو الموقع.",
                    "Dashboard صغيرة للحجوزات أو المبيعات أو المخزون.",
                    "متابعة تلقائية للـ leads اللي ما اتردش عليها.",
                    "Prototype لأداة داخلية للداتا أو التشغيل أو التقارير.",
                ],
                "outcome": "العميل يشوف الحل بإيده قبل ما يقرر يبني النسخة الكاملة.",
            },
            {
                "heading": "3. مساعد واتساب AI",
                "bullets": [
                    "يرد على الأسئلة المتكررة.",
                    "يجمع بيانات العميل ويأهل الـ lead.",
                    "يحجز مواعيد ويبعت reminders.",
                    "يتابع الطلبات أو الحالات.",
                    "يحوّل الحالات الصعبة لموظف.",
                ],
                "outcome": "مناسب جدا للعيادات، العقارات، المطاعم، الـ e-commerce، وأي شركة بيعها أو خدمتها ماشية على واتساب.",
            },
            {
                "heading": "4. نظام leads ومتابعة مبيعات",
                "bullets": [
                    "يوصل إعلانات Facebook وInstagram بالموقع وواتساب.",
                    "يدخل الـ leads في CRM أو Google Sheets.",
                    "يبعت follow-up تلقائي.",
                    "يدي score للـ lead ويوجهه للشخص الصح.",
                ],
                "outcome": "leads أقل بتضيع، رد أسرع، ورؤية أوضح لفريق المبيعات.",
            },
            {
                "heading": "5. خدمة عملاء AI",
                "bullets": [
                    "ردود أولية على واتساب، شات الموقع، الإيميل، والسوشيال.",
                    "تجميع البيانات الناقصة قبل ما الموظف يستلم.",
                    "تنظيم المحادثات والتصعيد للحالات المهمة.",
                ],
                "outcome": "خدمة أسرع من غير ما تزود ضغط على الفريق.",
            },
            {
                "heading": "6. Automation للتشغيل الداخلي",
                "bullets": [
                    "إدخال داتا، فواتير، تنبيهات مخزون، تتبع أوردرات.",
                    "تذكير مهام للموظفين وتقارير يومية أو أسبوعية.",
                    "تحديث CRM وDashboards داخلية.",
                ],
                "outcome": "شغل يدوي أقل، أخطاء أقل، وصورة أوضح لصاحب البيزنس.",
            },
            {
                "heading": "7. مواقع بيزنس جاهزة للتحويل",
                "bullets": [
                    "موبايل-first، عربي وإنجليزي، واتساب CTA، فورمات، SEO structure.",
                    "صفحات خدمات، حجز، Analytics، وربط بالـ CRM.",
                ],
                "outcome": "موقع يجيب inquiries حقيقية، مش brochure وخلاص.",
            },
            {
                "heading": "8. Content Engine للسوشيال",
                "bullets": [
                    "أفكار posts حسب الصناعة.",
                    "تقويم 30 يوم.",
                    "Captions عربي وإنجليزي بنفس voice البراند.",
                    "Campaign ideas وقوالب قابلة لإعادة الاستخدام.",
                ],
                "outcome": "محتوى أسرع وأكثر ثباتا من غير استنزاف الفريق.",
            },
            {
                "heading": "9. E-commerce optimization",
                "bullets": [
                    "وصف منتجات، recovery للعربيات المتروكة، review collection.",
                    "تنبيهات مخزون، support automation، وسلاسل follow-up.",
                ],
                "outcome": "مبيعات راجعة أكتر وتجربة شراء أنضف.",
            },
            {
                "heading": "10. حجز ذكي",
                "bullets": [
                    "حجز من واتساب، reminders، rescheduling، أسئلة intake.",
                    "تنسيق calendar للموظفين وتقليل no-shows.",
                ],
                "outcome": "مواعيد أقل بتفوت وفريق استقبال عليه ضغط أقل.",
            },
            {
                "heading": "11. Analytics dashboards",
                "bullets": [
                    "Leads، مبيعات، حجوزات، محادثات، support volume، مخزون، ROI.",
                    "الأرقام اللي صاحب البيزنس محتاج يشوفها كل يوم في مكان واحد.",
                ],
                "outcome": "قرارات على أرقام بدل التخمين.",
            },
            {
                "heading": "12. AI strategy وتنفيذ كامل",
                "bullets": [
                    "Workflow audit، roadmap، 3 إلى 5 فرص automation، implementation plan.",
                    "أدوات مخصوصة، onboarding للفريق، وتحسين شهري.",
                ],
                "outcome": "شريك AI طويل المدى يحسن التشغيل شهر ورا شهر.",
            },
        ],
    },
    {
        "title": "مين نتصل بيهم الأول؟",
        "body": [
            "ابدأ بأي بيزنس عنده حجم رسائل عالي، بيع أو حجز على واتساب، أسئلة متكررة، أو فريق صغير بيحاول يلحق خدمة العملاء والمبيعات.",
        ],
        "items": [
            {
                "heading": "أفضل قطاعات للاتصال",
                "bullets": [
                    "E-commerce وD2C brands: أوردرات، مقاسات، شحن، returns، ومتابعة عربيات متروكة.",
                    "Real estate: leads كتير، أسئلة عقارات متكررة، follow-up ضعيف، وCRM غير منتظم.",
                    "Medical clinics: حجوزات، reminders، أسئلة مرضى، no-shows، وضغط استقبال.",
                    "F&B وbakeries: أوردرات، delivery areas، custom orders، reviews، وكاترينج.",
                    "Beauty, dental, wellness: مواعيد، أسعار، consultation requests، وWhatsApp confirmation.",
                    "Furniture/custom products: مقاسات، خامات، صور reference، عروض أسعار، ومتابعة تصنيع.",
                ],
            }
        ],
    },
    {
        "title": "زاوية البيع",
        "items": [
            {
                "heading": "ابدأ بالألم، مش بالتكنولوجيا",
                "bullets": [
                    "ما تقولش: بنعمل AI automation.",
                    "قول: بنشيل الشغل المتكرر اللي بيضيع وقت فريقك كل يوم.",
                ],
            },
            {
                "heading": "بيع الديمو، مش المشروع",
                "bullets": [
                    "قبل ما نطلب منك تدفع، هنعمل ديمو صغير على مشكلة حقيقية عندك عشان تشوف الحل شغال.",
                    "لو الديمو مفيد، نتكلم في النسخة الكاملة. لو مش مناسب، تبقى خرجت بفكرة أوضح عن اللي ممكن يتحسن.",
                ],
            },
            {
                "heading": "ركّز على النتائج",
                "bullets": [
                    "ساعات أقل بتضيع.",
                    "رد فوري على الـ leads.",
                    "استفسارات أقل بتتفوت.",
                    "متابعة أحسن.",
                    "صورة أوضح لصاحب البيزنس.",
                ],
            },
        ],
    },
    {
        "title": "سكريبتات جاهزة",
        "items": [
            {
                "heading": "سكريبت مكالمة قصيرة",
                "body": [
                    "أهلا يا [الاسم]، معاك [اسمك] من Sierra Tech Spaces.",
                    "احنا بنساعد الشركات تشيل الشغل المتكرر وترد أسرع على العملاء وتمنع ضياع الـ leads.",
                    "لاحظنا إن بيزنس زي بتاعكم غالبا بيبقى عنده أسئلة متكررة، متابعة واتساب، حجوزات، أو شغل إداري بياخد وقت.",
                    "فكرتنا بسيطة: نسمع منك 15 دقيقة، نفهم أكبر حاجة مبطأة الشغل، وبعدين نعمل ديمو صغير مجاني يورّيك automation ممكن يشتغل ازاي عندك.",
                    "ينفع ناخد call سريع الأسبوع ده؟",
                ],
            },
            {
                "heading": "رسالة واتساب",
                "body": [
                    "أهلا يا [الاسم]، أنا [اسمك] من Sierra Tech Spaces.",
                    "احنا بنساعد البيزنسات توفر وقت في الردود، المتابعة، الحجوزات، التقارير، والمهام المتكررة.",
                    "بنبدأ بحاجة بسيطة: نفهم أكبر pain point عندكم، وبعدها نعمل ديمو مجاني مخصوص على workflow حقيقي من شغلكم.",
                    "لو الديمو مفيد، نتكلم في التنفيذ. لو مش مناسب، مفيش أي التزام.",
                    "ممكن 15 دقيقة call نشوف ايه ممكن يتعمله automation؟",
                ],
            },
            {
                "heading": "أسئلة discovery",
                "bullets": [
                    "ايه أكتر جزء في الشغل بياخد مجهود يدوي كل أسبوع؟",
                    "العملاء بيكلموكم فين أكتر: واتساب، Instagram، تليفون، موقع؟",
                    "هل في leads بتضيع عشان الرد أو المتابعة بتتأخر؟",
                    "ايه أكتر أسئلة بتتكرر من العملاء؟",
                    "لو شيلنا مهمة واحدة مؤلمة الشهر ده، تحب تبقى ايه؟",
                ],
            },
            {
                "heading": "قفل المكالمة",
                "body": [
                    "من كلامك، واضح إن أكبر فرصة عندكم هي [pain point].",
                    "بدل ما نشرح نظري، نقدر نعمل ديمو بسيط يورّي ازاي [solution] ممكن توفر وقت وتحسن الردود وتنظم المتابعة.",
                    "الديمو مجاني ومفيش أي ضغط. ابعتلنا [FAQs، screenshots، workflow، sample leads] واحنا نبني أول نسخة مصغرة.",
                ],
            },
            {
                "heading": "اعتراضات وردود",
                "bullets": [
                    "لسه مش جاهزين لـ AI: تمام. احنا مش بنبدأ بـ AI. بنبدأ بمشكلة متكررة ونشوف هل automation ينفع يحلها ولا لأ.",
                    "غالي؟ السعر حسب السكوب، بس مفيش التزام قبل ما تشوف ديمو واضح.",
                    "عندنا موظفين بيعملوا ده: الهدف مش نستبدل الفريق. الهدف نشيل التكرار عشان الفريق يركز على العملاء والبيع.",
                    "هيشتغل بالعربي؟ آه. بنبني عربي-first، وبالمصري لو محتاجين، خصوصا لواتساب وخدمة العملاء والحجز.",
                ],
            },
        ],
    },
]


PROSPECTS = [
    {
        "rank": 1,
        "priority": "A",
        "name": "Polymedica Clinics",
        "sector": "عيادات متعددة التخصصات - New Cairo",
        "why": "عندهم واتساب، تليفون، وحجز online، فالألم المتوقع هو ضغط استقبال وحجوزات وأسئلة متكررة.",
        "offer": "مساعد واتساب للحجز والأسئلة + reminders للمرضى.",
        "contact": "WhatsApp / phone موجودين على صفحة التواصل الرسمية",
        "source": "https://polymedicaclinics.com/contact",
    },
    {
        "rank": 2,
        "priority": "A",
        "name": "Zein Medical Center",
        "sector": "مركز طبي - التجمع الأول",
        "why": "موقعهم فيه online appointment وخدمات طبية كثيرة، وده مناسب لنظام حجز وتأهيل مرضى.",
        "offer": "Booking automation + patient intake questions + reminders.",
        "contact": "Phone / form من الموقع الرسمي",
        "source": "https://zeinmedical.com/en",
    },
    {
        "rank": 3,
        "priority": "A",
        "name": "H Dental Institute",
        "sector": "Dental clinic - Cairo",
        "why": "بيستخدموا WhatsApp وonline booking، وفيه مواعيد واستفسارات وأسعار وخدمات متعددة.",
        "offer": "WhatsApp receptionist + booking + follow-up بعد consultation.",
        "contact": "WhatsApp / booking من صفحة التواصل",
        "source": "https://hdentalinstitute.com/pages/contact",
    },
    {
        "rank": 4,
        "priority": "A",
        "name": "Revita Dental Clinic",
        "sector": "Dental clinic - New Cairo",
        "why": "عندهم phone وWhatsApp وخدمات كثيرة، مناسبين جدا لحجز ذكي وردود أولية.",
        "offer": "حجز واتساب + FAQ عن الخدمات + reminders.",
        "contact": "Phone / WhatsApp على الموقع",
        "source": "https://revitadentalclinic.com/",
    },
    {
        "rank": 5,
        "priority": "A",
        "name": "Luma Hair & Wellness Studio",
        "sector": "Salon / wellness - Zamalek",
        "why": "الحجز عندهم online وWhatsApp، ومواعيد الصالونات بتحتاج confirmations وتقليل no-shows.",
        "offer": "Booking flow + WhatsApp confirmation + rescheduling.",
        "contact": "Appointment form / WhatsApp من الموقع",
        "source": "https://lumahair.org/",
    },
    {
        "rank": 6,
        "priority": "A",
        "name": "RealEstate.eg",
        "sector": "Real estate platform",
        "why": "بيذكروا 75+ developers و500+ projects وWhatsApp رد سريع، فحجم leads عالي ومناسب للتأهيل والrouting.",
        "offer": "Lead qualification bot + CRM routing + dashboard.",
        "contact": "Call / WhatsApp / form من صفحة التواصل",
        "source": "https://realestate.eg/en/contact-us",
    },
    {
        "rank": 7,
        "priority": "A",
        "name": "Leverage Real Estate",
        "sector": "Real estate agency - New Cairo",
        "why": "عندهم consultation وWhatsApp ومتابعة عقارات، وده pain واضح في الreal estate.",
        "offer": "Lead intake + property matching questions + follow-up automation.",
        "contact": "WhatsApp / phone / form من الموقع",
        "source": "https://leverageg.com/",
    },
    {
        "rank": 8,
        "priority": "A",
        "name": "Platform Real Estate",
        "sector": "Real estate - New Cairo",
        "why": "صفحة التواصل فيها Call وWhatsApp، وده مناسب لتأهيل الـ leads قبل ما توصل للسيلز.",
        "offer": "WhatsApp qualification + sales handoff + lead sheet.",
        "contact": "Call / WhatsApp من صفحة التواصل",
        "source": "https://platformrealestate.co/en/contact/",
    },
    {
        "rank": 9,
        "priority": "B",
        "name": "Marker Group",
        "sector": "Real estate / reservations",
        "why": "الموقع بيطلب reservations وWhatsApp، فالمشكلة غالبا في الاستفسارات والمتابعة.",
        "offer": "Reservation inquiry assistant + lead tracking.",
        "contact": "WhatsApp / phone / email من صفحة التواصل",
        "source": "https://www.markergroup.net/Contact",
    },
    {
        "rank": 10,
        "priority": "A",
        "name": "Deqah Furniture",
        "sector": "Custom furniture - Damietta",
        "why": "العميل بيبعت مقاسات وصور وميزانية على واتساب، وده workflow ممتاز للأتمتة.",
        "offer": "WhatsApp quote intake + material questions + lead dashboard.",
        "contact": "Phone / WhatsApp من الموقع",
        "source": "https://deqahfurniture.com/",
    },
    {
        "rank": 11,
        "priority": "A",
        "name": "Daleenz",
        "sector": "Furniture retail / factory - Cairo",
        "why": "عندهم showroom وfactory وWhatsApp، ومناسبين لتجميع طلبات وعروض أسعار.",
        "offer": "Product inquiry bot + quote capture + showroom appointment flow.",
        "contact": "WhatsApp / phone / email من صفحة التواصل",
        "source": "https://daleenz.com/pages/contact-us",
    },
    {
        "rank": 12,
        "priority": "B",
        "name": "Asal Furniture",
        "sector": "Furniture - Damietta / Cairo branches",
        "why": "عندهم فروع كثيرة وWhatsApp في صفحة التواصل، وده مناسب لتوحيد inquiries.",
        "offer": "Branch routing + custom order intake + follow-up.",
        "contact": "Phones / WhatsApp / emails من صفحة التواصل",
        "source": "https://asalfurniture.com/contact/",
    },
    {
        "rank": 13,
        "priority": "A",
        "name": "Sigma Fit",
        "sector": "Sportswear e-commerce",
        "why": "صفحة التواصل فيها WhatsApp وMessenger وInstagram، والشحن والreturns بيخلقوا أسئلة متكررة.",
        "offer": "E-commerce support bot + order FAQ + abandoned cart follow-up.",
        "contact": "WhatsApp / Messenger / Instagram / email من الموقع",
        "source": "https://sigmafiteg.com/pages/contact-us",
    },
    {
        "rank": 14,
        "priority": "B",
        "name": "Dresscode",
        "sector": "Fashion e-commerce",
        "why": "بيشحنوا لكل مصر وعندهم returns، فمناسبين لخدمة عملاء وردود مقاسات وشحن.",
        "offer": "Customer support automation + returns FAQ + product questions.",
        "contact": "Phone / email / contact form من الموقع",
        "source": "https://dresscodeme.com/pages/contact",
    },
    {
        "rank": 15,
        "priority": "B",
        "name": "Dressentials",
        "sector": "Fashion e-commerce",
        "why": "لديهم WhatsApp مباشر، ومناسبين لردود المقاسات والتوفر والطلبات.",
        "offer": "WhatsApp shopping assistant + order questions + follow-up.",
        "contact": "WhatsApp / email من صفحة التواصل",
        "source": "https://dressentials.co/pages/contact",
    },
    {
        "rank": 16,
        "priority": "A",
        "name": "Konouz Dessert Shop",
        "sector": "Bakery / desserts / catering",
        "why": "عندهم WhatsApp ordering، custom cakes، catering، وتوصيل، فالأتمتة هتوفر وقت بسرعة.",
        "offer": "Order-taking WhatsApp bot + delivery zones + custom cake intake.",
        "contact": "WhatsApp / phone / email من الموقع",
        "source": "https://konouz-dessert.com/",
    },
    {
        "rank": 17,
        "priority": "B",
        "name": "Gourmand Bakery",
        "sector": "Gluten-free bakery - Madinaty",
        "why": "يعتمدوا على WhatsApp delivery ومواعيد opening، مناسبين للطلبات المسبقة والمتابعة.",
        "offer": "Pre-order WhatsApp flow + delivery area routing + launch list follow-up.",
        "contact": "WhatsApp / email من الموقع",
        "source": "https://thegourmandbakery.com/",
    },
    {
        "rank": 18,
        "priority": "A",
        "name": "ZAT Aesthetic Clinics",
        "sector": "Aesthetic clinics - Madinaty / Rehab",
        "why": "الخدمات التجميلية فيها أسئلة أسعار وحجز واستفسارات متكررة قبل الزيارة.",
        "offer": "Consultation intake + booking + WhatsApp FAQ.",
        "contact": "WhatsApp / contact من الموقع",
        "source": "https://zat-clinic.com/",
    },
]


def add_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, size: int | None = None, bold: bool = False, color: RGBColor | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_bidi(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    run.font.rtl = True
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return paragraph


def add_docx_bullets(doc: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        add_docx_paragraph(doc, f"• {bullet}", size=10)


def build_markdown() -> str:
    lines: list[str] = [
        "# Sierra Tech Spaces - عرض العملاء وسكريبت البيع بالمصري",
        "",
        "نسخة Egyptian Arabic مبنية على العرض الإنجليزي الموجود في CUSTOMER_OFFER_AND_SCRIPT.md.",
        "",
    ]
    for section in OFFER_SECTIONS:
        lines.extend([f"## {section['title']}", ""])
        for paragraph in section.get("body", []):
            lines.extend([paragraph, ""])
        if "quote" in section:
            lines.extend([f"> {section['quote']}", ""])
        for item in section.get("items", []):
            lines.extend([f"### {item['heading']}", ""])
            for paragraph in item.get("body", []):
                lines.extend([paragraph, ""])
            for bullet in item.get("bullets", []):
                lines.append(f"- {bullet}")
            if item.get("bullets"):
                lines.append("")
            if "outcome" in item:
                lines.extend([f"**الناتج للعميل:** {item['outcome']}", ""])
    lines.extend(["## قائمة العملاء المحتملين للاتصال", ""])
    lines.extend(["الأولوية A معناها اتصل الأول. الأسماء دي مختارة لأنها ظاهرة في قنوات بيع أو حجز أو خدمة عملاء محتاجة رد سريع ومتابعة.", ""])
    lines.append("| # | الأولوية | العميل | القطاع | أول عرض نقوله | قناة التواصل |")
    lines.append("|---|---|---|---|---|---|")
    for p in PROSPECTS:
        lines.append(f"| {p['rank']} | {p['priority']} | {p['name']} | {p['sector']} | {p['offer']} | {p['contact']} |")
    lines.extend(["", "## مصادر قائمة الاتصال", ""])
    for p in PROSPECTS:
        lines.append(f"- {p['name']}: {p['source']}")
    lines.append("")
    return "\n".join(lines)


def build_prospect_markdown() -> str:
    lines = [
        "# Most Potential Customers To Call - Egypt",
        "",
        "Prioritized for Sierra Tech Spaces based on visible WhatsApp, booking, ordering, contact-form, or lead-heavy workflows. Verify contact details before outreach.",
        "",
        "| Rank | Priority | Customer | Sector | Why call | First offer | Contact channel | Source |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for p in PROSPECTS:
        lines.append(
            f"| {p['rank']} | {p['priority']} | {p['name']} | {p['sector']} | {p['why']} | {p['offer']} | {p['contact']} | {p['source']} |"
        )
    lines.append("")
    return "\n".join(lines)


def build_docx() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Tahoma"
        style._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.0))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bidi(title)
    run = title.add_run("Sierra Tech Spaces\nعرض العملاء وسكريبت البيع بالمصري")
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 20, 38)
    run.font.rtl = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bidi(subtitle)
    s_run = subtitle.add_run("اكتشاف المشكلة، ديمو مجاني، عروض خدمات، سكريبتات، وقائمة عملاء للاتصال.")
    s_run.font.name = "Tahoma"
    s_run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    s_run.font.size = Pt(10)
    s_run.font.color.rgb = RGBColor(100, 116, 139)
    s_run.font.rtl = True

    for section_data in OFFER_SECTIONS:
        add_docx_paragraph(doc, section_data["title"], size=15, bold=True, color=RGBColor(11, 20, 38))
        for paragraph in section_data.get("body", []):
            add_docx_paragraph(doc, paragraph, size=10)
        if "quote" in section_data:
            q = add_docx_paragraph(doc, section_data["quote"], size=10, bold=True, color=RGBColor(11, 20, 38))
            q.paragraph_format.left_indent = Cm(0.4)
            q.paragraph_format.right_indent = Cm(0.4)
        for item in section_data.get("items", []):
            add_docx_paragraph(doc, item["heading"], size=12, bold=True, color=RGBColor(6, 182, 212))
            for paragraph in item.get("body", []):
                add_docx_paragraph(doc, paragraph, size=10)
            add_docx_bullets(doc, item.get("bullets", []))
            if "outcome" in item:
                add_docx_paragraph(doc, f"الناتج للعميل: {item['outcome']}", size=10, bold=True)

    doc.add_page_break()
    add_docx_paragraph(doc, "قائمة العملاء المحتملين للاتصال", size=15, bold=True, color=RGBColor(11, 20, 38))
    add_docx_paragraph(
        doc,
        "الأولوية A اتصل بيها الأول. الاختيار مبني على وجود واتساب أو حجز أو طلبات أو leads كتير، وده يخلي الديمو المجاني سهل ومقنع.",
        size=10,
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["#", "أولوية", "العميل", "القطاع", "أول عرض", "قناة التواصل"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "0B1426")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_bidi(p)
        run = p.add_run(header)
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        run.font.rtl = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(8)

    for prospect in PROSPECTS:
        cells = table.add_row().cells
        values = [
            str(prospect["rank"]),
            prospect["priority"],
            prospect["name"],
            prospect["sector"],
            prospect["offer"],
            prospect["contact"],
        ]
        for cell, value in zip(cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_bidi(p)
            run = p.add_run(value)
            run.font.name = "Tahoma"
            run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
            run.font.rtl = True
            run.font.size = Pt(7.5)

    doc.add_paragraph()
    add_docx_paragraph(doc, "ملاحظة: اتأكد من أرقام وقنوات التواصل من الموقع الرسمي قبل أي outreach.", size=8, color=RGBColor(100, 116, 139))
    doc.save(DOCX_OUTPUT)


def shape(text: str) -> str:
    if not text:
        return ""
    return get_display(arabic_reshaper.reshape(text))


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return html.unescape(text)


class PdfWriter:
    def __init__(self, path: Path):
        PDF_DIR.mkdir(parents=True, exist_ok=True)
        pdfmetrics.registerFont(TTFont("Tahoma", FONT_REG))
        pdfmetrics.registerFont(TTFont("Tahoma-Bold", FONT_BOLD))
        self.path = path
        self.c = canvas.Canvas(str(path), pagesize=A4)
        self.width, self.height = A4
        self.left = 42
        self.right = self.width - 42
        self.top = self.height - 56
        self.bottom = 54
        self.y = self.top
        self.page = 0
        self.new_page(first=True)

    def new_page(self, first: bool = False) -> None:
        if not first:
            self.footer()
            self.c.showPage()
        self.page += 1
        self.y = self.top
        self.header()

    def header(self) -> None:
        self.c.setFillColor(NAVY)
        self.c.rect(0, self.height - 22, self.width, 22, stroke=0, fill=1)
        self.c.setFillColor(TEAL)
        self.c.rect(0, self.height - 22, 116, 22, stroke=0, fill=1)
        self.c.setFillColor(colors.white)
        self.c.setFont("Tahoma-Bold", 7.5)
        self.c.drawString(self.left, self.height - 15, "Sierra Tech Spaces")

    def footer(self) -> None:
        self.c.setFillColor(MUTED)
        self.c.setFont("Tahoma", 7)
        self.c.drawString(self.left, 28, "AI that works as hard as you do.")
        self.c.drawRightString(self.right, 28, f"Page {self.page}")

    def ensure_space(self, amount: float) -> None:
        if self.y - amount < self.bottom:
            self.new_page()

    def draw_rule(self) -> None:
        self.ensure_space(14)
        self.c.setStrokeColor(RULE)
        self.c.setLineWidth(0.5)
        self.c.line(self.left, self.y, self.right, self.y)
        self.y -= 12

    def wrap(self, text: str, font: str, size: float, max_width: float) -> list[str]:
        words = text.split()
        if not words:
            return []
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = word if not current else current + " " + word
            if pdfmetrics.stringWidth(shape(candidate), font, size) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines

    def para(
        self,
        text: str,
        *,
        font: str = "Tahoma",
        size: float = 9.4,
        leading: float = 14,
        color=SLATE,
        before: float = 0,
        after: float = 5,
        max_width: float | None = None,
    ) -> None:
        clean = strip_md(text)
        max_width = max_width or (self.right - self.left)
        lines = self.wrap(clean, font, size, max_width)
        self.ensure_space(before + max(after, len(lines) * leading))
        self.y -= before
        self.c.setFillColor(color)
        self.c.setFont(font, size)
        for line in lines:
            self.c.drawRightString(self.right, self.y, shape(line))
            self.y -= leading
        self.y -= after

    def bullet(self, text: str) -> None:
        self.para("• " + text, size=8.9, leading=13, after=2)

    def title(self, text: str, subtitle: str) -> None:
        if LOGO.exists():
            self.c.drawImage(str(LOGO), self.width / 2 - 34, self.y - 68, width=68, height=68, mask="auto")
            self.y -= 82
        self.para(text, font="Tahoma-Bold", size=20, leading=26, color=NAVY, after=5)
        self.para(subtitle, size=8.8, leading=12.5, color=MUTED, after=9)
        self.c.setStrokeColor(TEAL)
        self.c.setLineWidth(2)
        self.c.line(self.width / 2, self.y, self.right, self.y)
        self.c.setStrokeColor(GOLD)
        self.c.line(self.left, self.y, self.width / 2 - 8, self.y)
        self.y -= 18

    def prospect_card(self, p: dict) -> None:
        block_height = 70
        self.ensure_space(block_height)
        self.c.setFillColor(LIGHT_BG)
        self.c.roundRect(self.left, self.y - block_height + 8, self.right - self.left, block_height, 5, stroke=0, fill=1)
        self.y -= 8
        self.para(f"{p['rank']}. {p['name']} - أولوية {p['priority']}", font="Tahoma-Bold", size=10.2, leading=13, color=NAVY, after=1, max_width=self.right - self.left - 16)
        old_right = self.right
        self.right -= 8
        self.para(f"القطاع: {p['sector']}", size=8, leading=10, after=1, max_width=old_right - self.left - 18)
        self.para(f"ليه نتصل: {p['why']}", size=8, leading=10, after=1, max_width=old_right - self.left - 18)
        self.para(f"أول عرض: {p['offer']}", size=8, leading=10, after=1, max_width=old_right - self.left - 18)
        self.para(f"القناة: {p['contact']}", size=7.5, leading=9.5, color=MUTED, after=3, max_width=old_right - self.left - 18)
        self.right = old_right

    def save(self) -> None:
        self.footer()
        self.c.save()


def build_pdf() -> None:
    pdf = PdfWriter(PDF_OUTPUT)
    pdf.title(
        "Sierra Tech Spaces - عرض العملاء وسكريبت البيع بالمصري",
        "نسخة عملية للمكالمات والواتساب: اكتشاف المشكلة، ديمو مجاني، وعملاء محتملين للاتصال.",
    )
    for section in OFFER_SECTIONS:
        pdf.para(section["title"], font="Tahoma-Bold", size=14, leading=18, color=NAVY, before=5, after=4)
        for paragraph in section.get("body", []):
            pdf.para(paragraph)
        if "quote" in section:
            pdf.para(section["quote"], font="Tahoma-Bold", size=9.4, leading=13.5, color=NAVY, after=7)
        for item in section.get("items", []):
            pdf.para(item["heading"], font="Tahoma-Bold", size=10.5, leading=13, color=TEAL, before=3, after=2)
            for paragraph in item.get("body", []):
                pdf.para(paragraph)
            for bullet in item.get("bullets", []):
                pdf.bullet(bullet)
            if "outcome" in item:
                pdf.para(f"الناتج للعميل: {item['outcome']}", font="Tahoma-Bold", size=8.8, leading=12.5, color=NAVY, after=4)
        pdf.draw_rule()

    pdf.new_page()
    pdf.para("قائمة العملاء المحتملين للاتصال", font="Tahoma-Bold", size=15, leading=19, color=NAVY, after=4)
    pdf.para(
        "اتصل بالأولوية A الأول. كل اسم هنا عنده مؤشر واضح إن عنده رسائل أو حجوزات أو طلبات أو leads محتاجة رد سريع ومتابعة.",
        size=9,
        leading=13,
        color=SLATE,
    )
    for prospect in PROSPECTS:
        pdf.prospect_card(prospect)
    pdf.save()


def write_csv() -> None:
    with PROSPECT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["rank", "priority", "name", "sector", "why", "offer", "contact", "source"],
        )
        writer.writeheader()
        writer.writerows(PROSPECTS)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_MD.write_text(build_markdown(), encoding="utf-8")
    PROSPECT_MD.write_text(build_prospect_markdown(), encoding="utf-8")
    write_csv()
    build_docx()
    build_pdf()
    print(DOCX_OUTPUT)
    print(PDF_OUTPUT)
    print(PROSPECT_MD)
    print(PROSPECT_CSV)


if __name__ == "__main__":
    main()
